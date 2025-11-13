#!/usr/bin/env python3
# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx>=1.1.0",
# ]
# ///

"""
python-docx demonstration: Creating and manipulating Word documents

This script demonstrates key features of the python-docx library including:
- Creating documents
- Adding paragraphs with various styles
- Adding headings
- Formatting text (bold, italic, underline)
- Creating tables
- Adding lists
- Working with sections and page layout
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    """Demonstrate python-docx capabilities."""
    print("=" * 70)
    print("python-docx Demonstration: Creating and Manipulating Word Documents")
    print("=" * 70)
    print()

    # Line 36: Create a new Document object
    print("[Line 36] Creating a new Word document...")
    document = Document()
    print("✓ Document created successfully")
    print()

    # Line 42: Add a title heading
    print("[Line 42] Adding a title heading...")
    title = document.add_heading("python-docx Feature Demonstration", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("✓ Title added: 'python-docx Feature Demonstration'")
    print()

    # Line 49: Add a regular paragraph with text formatting
    print("[Line 49] Adding a paragraph with formatted text...")
    intro = document.add_paragraph("This document demonstrates the ")
    intro.add_run("python-docx").bold = True
    intro.add_run(" library for creating and manipulating Word documents. ")
    intro.add_run("It's powerful and easy to use!").italic = True
    print("✓ Paragraph added with bold and italic formatting")
    print()

    # Line 58: Add a section heading
    print("[Line 58] Adding section heading: 'Basic Text Formatting'...")
    document.add_heading("1. Basic Text Formatting", level=1)
    print("✓ Section heading added")
    print()

    # Line 64: Demonstrate various text formatting options
    print("[Line 64] Adding paragraph with various formatting options...")
    p = document.add_paragraph()
    p.add_run("Bold text").bold = True
    p.add_run(", ")
    p.add_run("Italic text").italic = True
    p.add_run(", ")
    p.add_run("Underlined text").underline = True
    p.add_run(", and ")
    colored_run = p.add_run("Colored text")
    colored_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    colored_run.font.size = Pt(14)
    print("✓ Formatted text added: bold, italic, underline, color, and size")
    print()

    # Line 79: Add a heading for lists section
    print("[Line 79] Adding section heading: 'Lists'...")
    document.add_heading("2. Lists", level=1)
    print("✓ Section heading added")
    print()

    # Line 85: Create a bulleted list
    print("[Line 85] Creating a bulleted list...")
    document.add_paragraph("Features of python-docx:")
    document.add_paragraph("Create new documents", style="List Bullet")
    document.add_paragraph("Add headings and paragraphs", style="List Bullet")
    document.add_paragraph("Format text (bold, italic, etc.)", style="List Bullet")
    document.add_paragraph("Create tables", style="List Bullet")
    document.add_paragraph("Insert images", style="List Bullet")
    print("✓ Bulleted list created with 5 items")
    print()

    # Line 96: Create a numbered list
    print("[Line 96] Creating a numbered list...")
    document.add_paragraph("Steps to use python-docx:")
    document.add_paragraph("Install: pip install python-docx", style="List Number")
    document.add_paragraph("Import: from docx import Document", style="List Number")
    document.add_paragraph("Create: document = Document()", style="List Number")
    document.add_paragraph("Add content using various methods", style="List Number")
    document.add_paragraph("Save: document.save('filename.docx')", style="List Number")
    print("✓ Numbered list created with 5 steps")
    print()

    # Line 107: Add a heading for tables section
    print("[Line 107] Adding section heading: 'Tables'...")
    document.add_heading("3. Tables", level=1)
    print("✓ Section heading added")
    print()

    # Line 113: Create a table
    print("[Line 113] Creating a 4x3 table...")
    table = document.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    print("✓ Table created with 'Light Grid Accent 1' style")
    print()

    # Line 120: Populate table headers
    print("[Line 120] Populating table headers...")
    header_cells = table.rows[0].cells
    header_cells[0].text = "Library"
    header_cells[1].text = "Purpose"
    header_cells[2].text = "Difficulty"
    print("✓ Headers: Library, Purpose, Difficulty")
    print()

    # Line 129: Populate table data
    print("[Line 129] Populating table data...")
    data_rows = [
        ("python-docx", "Word documents", "Easy"),
        ("openpyxl", "Excel spreadsheets", "Easy"),
        ("PyPDF2", "PDF manipulation", "Medium"),
    ]

    for i, (lib, purpose, difficulty) in enumerate(data_rows, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = lib
        row_cells[1].text = purpose
        row_cells[2].text = difficulty

    print(f"✓ Added {len(data_rows)} data rows to table")
    print()

    # Line 147: Add a heading for page layout section
    print("[Line 147] Adding section heading: 'Page Layout'...")
    document.add_heading("4. Page Layout and Sections", level=1)
    print("✓ Section heading added")
    print()

    # Line 153: Demonstrate section manipulation
    print("[Line 153] Adding information about current page layout...")
    section = document.sections[0]
    p = document.add_paragraph(f"Page width: {section.page_width.inches:.2f} inches")
    p.add_run(f"\nPage height: {section.page_height.inches:.2f} inches")
    p.add_run(f"\nTop margin: {section.top_margin.inches:.2f} inches")
    p.add_run(f"\nBottom margin: {section.bottom_margin.inches:.2f} inches")
    p.add_run(f"\nLeft margin: {section.left_margin.inches:.2f} inches")
    p.add_run(f"\nRight margin: {section.right_margin.inches:.2f} inches")
    print("✓ Page layout information added")
    print()

    # Line 166: Add a conclusion heading
    print("[Line 166] Adding conclusion section...")
    document.add_heading("5. Conclusion", level=1)
    document.add_paragraph(
        "The python-docx library provides a comprehensive and intuitive API for "
        "creating and manipulating Word documents programmatically. It supports "
        "all common document elements and formatting options."
    )
    print("✓ Conclusion paragraph added")
    print()

    # Line 177: Save the document
    output_file = "python_docx_demo.docx"
    print(f"[Line 177] Saving document to '{output_file}'...")
    document.save(output_file)
    print(f"✓ Document saved successfully to '{output_file}'")
    print()

    # Line 184: Display summary statistics
    print("=" * 70)
    print("Document Creation Summary")
    print("=" * 70)
    print(f"Total paragraphs: {len(document.paragraphs)}")
    print(f"Total tables: {len(document.tables)}")
    print(f"Total sections: {len(document.sections)}")
    print(f"Output file: {output_file}")
    print()
    print("✓ Document creation complete!")
    print("=" * 70)


if __name__ == "__main__":
    main()
